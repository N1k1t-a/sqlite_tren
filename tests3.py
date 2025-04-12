from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Создаем новый документ с учетом правил оформления и выравниванием нумерации страниц по центру



doc = Document()

# Устанавливаем поля
section = doc.sections[0]
section.left_margin = Cm(3)  # Левое поле 30 мм
section.right_margin = Cm(1.5)  # Правое поле 15 мм
section.top_margin = Cm(2)  # Верхнее поле 20 мм
section.bottom_margin = Cm(2)  # Нижнее поле 20 мм

# Устанавливаем шрифт и абзацный отступ
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)  # Шрифт 12 пт
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.5  # Межстрочный интервал 1,5
paragraph_format.first_line_indent = Cm(1.25)  # Абзацный отступ 1,25 см

# Титульный лист
doc.add_paragraph("\n" * 10)
doc.add_paragraph("Государственный университет аэрокосмического приборостроения", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Отчет по лабораторной работе", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("\n" * 5)
doc.add_paragraph("Студент: Иванов И.И.", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Группа: 123-45", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Преподаватель: Петров П.П.", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("\n" * 10)
doc.add_paragraph("Санкт-Петербург 2024", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Разрыв страницы после титульного листа
doc.add_page_break()

# Содержание
doc.add_paragraph("Содержание", style='Heading 1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("\n")
doc.add_paragraph("Введение .................................................................................. 2")
doc.add_paragraph("1. Заголовок 1-го уровня ............................................................... 3")
doc.add_paragraph("1.1 Заголовок 2-го уровня ........................................................ 4")
doc.add_paragraph("Заключение ............................................................................. 5")
doc.add_paragraph("Список источников ................................................................. 6")

# Разрыв страницы после содержания
doc.add_page_break()

# Введение
doc.add_paragraph("Введение", style='Heading 1').bold = True
doc.add_paragraph("В данном отчете рассматриваются основные принципы оформления документов в соответствии с требованиями ГОСТ.")

# Заголовки 1-го и 2-го уровня
doc.add_paragraph("1. Заголовок 1-го уровня", style='Heading 1').bold = True
doc.add_paragraph("1.1 Заголовок 2-го уровня", style='Heading 2').bold = True

# Основной текст
doc.add_paragraph("Это основной текст отчета, где будут продемонстрированы различные элементы, такие как рисунки, таблицы и формулы.")

# Добавление рисунка с подписью
doc.add_paragraph("\nРисунок 1 - Пример изображения\n").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Добавление таблицы с подписью
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
doc.add_paragraph("\nТаблица 1 - Пример таблицы\n").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Формула
doc.add_paragraph("Пример формулы:").bold = True
doc.add_paragraph("E = mc^2")

# Заключение
doc.add_paragraph("Заключение", style='Heading 1').bold = True
doc.add_paragraph("В данном отчете был продемонстрирован процесс оформления отчетов по требованиям ГОСТ, включая все основные элементы.")

# Список источников
doc.add_page_break()
doc.add_paragraph("Список источников", style='Heading 1').bold = True
doc.add_paragraph("1. ГОСТ 7.32-2017. Система стандартов по информации, библиотечному и издательскому делу.")

# Нумерация страниц по центру снизу
footer = section.footer
paragraph = footer.paragraphs[0]
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
field_code = OxmlElement('w:fldSimple')
field_code.set(qn('w:instr'), 'PAGE')
run = paragraph.add_run()
run._r.append(field_code)

# Сохранение документа
file_path = r"C:\Users\aboob\Downloads\ГОСТ_отчет_пример(1).docx"
doc.save(file_path)

file_path



